"""
导出服务：将小说章节导出为 .docx（Microsoft Word）文件。

排版规范
--------
- 封面页：项目名称居中，Heading 1 样式
- 章节标题：Heading 1，黑体 16pt，加粗
- 正文：宋体（SimSun），12pt，1.5 倍行距
- 章节间分页
"""

from __future__ import annotations

import re
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ---- 章节正文字体设置：宋体（SimSun）----
_BODY_CN_FONT = "SimSun"          # 中文回退
_BODY_Western_FONT = "Times New Roman"  # 英文/数字回退


def _set_run_font(run, cn_font: str, size_pt: float) -> None:
    """设置 run 的中文字体、西文字体和字号。"""
    run.font.size = Pt(size_pt)
    run.font.name = _BODY_Western_FONT
    # rFonts.rPr 才能控制中文东亚字体
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:eastAsia"), cn_font)
    rfonts.set(qn("w:ascii"), _BODY_Western_FONT)
    rfonts.set(qn("w:hAnsi"), _BODY_Western_FONT)


def _set_paragraph_line_spacing(paragraph, line_spacing: float = 1.5) -> None:
    """设置段落的固定行距（倍数）。"""
    ppr = paragraph._element.get_or_add_pPr()
    spacing = ppr.find(qn("w:spacing"))
    if spacing is None:
        spacing = ppr.makeelement(qn("w:spacing"), {})
        ppr.append(spacing)
    spacing.set(qn("w:lineRule"), "auto")
    spacing.set(qn("w:line"), str(int(line_spacing * 240)))


def _extract_text(content: object) -> str:
    """从 chapter.content（JSON）中提取纯文本正文。"""
    if content is None:
        return ""
    if isinstance(content, dict):
        return (content.get("text") or content.get("content") or "").strip()
    if isinstance(content, str):
        return content.strip()
    import json

    return json.dumps(content, ensure_ascii=False)


def _sanitize_filename(name: str) -> str:
    """清理文件名中的非法字符。"""
    return re.sub(r'[\\/:*?"<>|\r\n\t]+', "_", name).strip()


def _build_heading1(doc: Document, title: str) -> None:
    """新增 Heading 1 标题段：黑体 16pt 加粗。"""
    heading = doc.add_heading(level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_paragraph_line_spacing(heading, 1.5)
    run = heading.add_run(title)
    _set_run_font(run, "SimHei", 16.0)
    run.bold = True


def _build_cover(doc: Document, project_name: str) -> None:
    """封面页：项目名称居中，Heading 1 样式。"""
    # 空行，使标题居中靠下
    for _ in range(6):
        p = doc.add_paragraph()
        _set_paragraph_line_spacing(p, 1.5)

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.style = doc.styles["Heading 1"]
    _set_paragraph_line_spacing(title_p, 1.5)
    run = title_p.add_run(project_name)
    _set_run_font(run, "SimHei", 24.0)
    run.bold = True

    doc.add_page_break()


def _build_chapter(doc: Document, chapter_number: int, title: str, text: str) -> None:
    """写一个章节：标题 + 正文。若已有内容则在章节前分页。"""
    if doc.paragraphs:
        doc.add_page_break()

    _build_heading1(doc, f"第{chapter_number}章  {title}")

    if not text:
        placeholder = doc.add_paragraph("（本章暂无正文）")
        _set_paragraph_line_spacing(placeholder, 1.5)
        for run in placeholder.runs:
            _set_run_font(run, _BODY_CN_FONT, 12.0)
        return

    # 按空行切段，保持换行结构
    for para_text in re.split(r"\n\s*\n", text):
        para_text = para_text.strip()
        if not para_text:
            continue
        p = doc.add_paragraph()
        _set_paragraph_line_spacing(p, 1.5)
        run = p.add_run(para_text)
        _set_run_font(run, _BODY_CN_FONT, 12.0)
        p.paragraph_format.first_line_indent = Pt(24)  # 2 字符首行缩进


def build_docx_bytes(
    project_name: str,
    chapters: Iterable[object],
) -> bytes:
    """
    生成 .docx 字节流。

    chapters：可迭代对象，每个元素需具备属性：
        chapter_number (int), title (str), content (JSON)
    """
    doc = Document()

    # 设置全局默认字体（正文样式），兜底
    normal = doc.styles["Normal"]
    normal.font.size = Pt(12)
    normal.font.name = _BODY_Western_FONT
    rpr = normal._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:eastAsia"), _BODY_CN_FONT)
    rfonts.set(qn("w:ascii"), _BODY_Western_FONT)
    rfonts.set(qn("w:hAnsi"), _BODY_Western_FONT)
    ppr = normal._element.get_or_add_pPr()
    spacing = ppr.find(qn("w:spacing"))
    if spacing is None:
        spacing = ppr.makeelement(qn("w:spacing"), {})
        ppr.append(spacing)
    spacing.set(qn("w:lineRule"), "auto")
    spacing.set(qn("w:line"), str(int(1.5 * 240)))

    safe_name = _sanitize_filename(project_name)

    # 封面
    _build_cover(doc, safe_name)

    # 章节
    for ch in chapters:
        ch_num = getattr(ch, "chapter_number", None)
        ch_title = getattr(ch, "title") or "未命名"
        text = _extract_text(getattr(ch, "content", None))
        _build_chapter(doc, ch_num or 0, ch_title, text)

    buf = Path("/tmp") / f"{safe_name}.docx"
    doc.save(str(buf))
    return buf.read_bytes()
